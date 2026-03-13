#!/usr/bin/env python3
import os
import re
import csv
import io
import requests
import feedparser
from datetime import datetime
from docx import Document

DESKTOP_DOC = '/Users/taofangzheng/Desktop/美股交易日复盘-持续更新.docx'

INDEX_SYMBOLS = {
    '^DJI': '道琼斯工业指数',
    '^SPX': '标普500指数',
    '^NDQ': '纳斯达克综合指数',
}

RSS_SOURCES = [
    ('CNBC Markets', 'https://www.cnbc.com/id/15839135/device/rss/rss.html'),
    ('MarketWatch Top Stories', 'https://www.marketwatch.com/rss/topstories'),
]

KEYWORDS = [
    'stocks', 'stock', 'dow', 's&p', 'nasdaq', 'fed', 'treasury', 'yield', 'inflation',
    'tariff', 'oil', 'market', 'economy', 'recession', 'rates', 'earnings', 'trump', 'china'
]


def fetch_csv(url: str) -> str:
    r = requests.get(url, timeout=30)
    r.raise_for_status()
    return r.text


def get_latest_two_days(symbol: str):
    url = f'https://stooq.com/q/d/l/?s={requests.utils.quote(symbol)}&i=d'
    text = fetch_csv(url)
    rows = list(csv.DictReader(io.StringIO(text)))
    rows = [r for r in rows if r.get('Close') and r['Close'] != 'N/D']
    if len(rows) < 2:
        return None, None
    return rows[-2], rows[-1]


def pct_change(prev_close: float, close: float) -> float:
    return (close - prev_close) / prev_close * 100


def fetch_headlines(limit=12):
    out = []
    for source_name, url in RSS_SOURCES:
        feed = feedparser.parse(url)
        for e in feed.entries:
            title = (e.get('title') or '').strip()
            link = (e.get('link') or '').strip()
            summary = re.sub(r'<[^>]+>', '', (e.get('summary') or '')).strip()
            text = f"{title} {summary}".lower()
            if any(k in text for k in KEYWORDS):
                out.append({
                    'source': source_name,
                    'title': title,
                    'link': link,
                })
    # 去重
    seen = set()
    uniq = []
    for x in out:
        key = x['title']
        if key in seen:
            continue
        seen.add(key)
        uniq.append(x)
    return uniq[:limit]


def build_key_factors(index_moves, headlines):
    factors = []
    # 价格行为
    sorted_moves = sorted(index_moves.items(), key=lambda kv: kv[1]['pct'])
    worst = sorted_moves[0]
    best = sorted_moves[-1]
    factors.append(f"指数层面：{worst[1]['name']}跌幅居前（{worst[1]['pct']:.2f}%），{best[1]['name']}相对抗跌/领涨（{best[1]['pct']:.2f}%），显示市场存在结构性分化。")

    # 新闻关键词归纳
    joined = ' '.join([h['title'].lower() for h in headlines])
    if any(k in joined for k in ['fed', 'rates', 'yield', 'treasury', 'inflation']):
        factors.append('利率与宏观预期：美联储路径、通胀和美债收益率预期仍是定价核心，压制高估值板块弹性。')
    if any(k in joined for k in ['tariff', 'china', 'trump', 'trade']):
        factors.append('政策与地缘不确定性：关税/贸易/地缘线索抬升风险溢价，导致资金更偏防御。')
    if any(k in joined for k in ['oil', 'energy']):
        factors.append('大宗商品波动：油价变化通过通胀预期与企业成本端，间接影响权益估值。')
    if any(k in joined for k in ['earnings', 'guidance', 'profit']):
        factors.append('财报与指引：业绩超预期但指引保守，或指引下修，都会放大个股与行业分化。')

    if len(factors) < 4:
        factors.append('资金面与风险偏好：在缺乏强增量催化时，市场更依赖仓位调整与短线情绪驱动。')

    return factors[:5]


def ensure_doc(path):
    if os.path.exists(path):
        return Document(path)
    doc = Document()
    doc.add_heading('美股交易日复盘（持续更新）', level=0)
    doc.add_paragraph('说明：本文件按交易日持续追加。内容包含指数表现、新闻要点与影响走势因素。')
    return doc


def main():
    now = datetime.now()
    doc = ensure_doc(DESKTOP_DOC)

    # 指数数据
    index_moves = {}
    trade_date = None
    for symbol, name in INDEX_SYMBOLS.items():
        prev_row, last_row = get_latest_two_days(symbol)
        if not prev_row or not last_row:
            continue
        prev_close = float(prev_row['Close'])
        close = float(last_row['Close'])
        chg = close - prev_close
        pct = pct_change(prev_close, close)
        trade_date = last_row['Date']
        index_moves[symbol] = {
            'name': name,
            'close': close,
            'chg': chg,
            'pct': pct,
            'date': last_row['Date'],
        }

    if not index_moves:
        raise RuntimeError('未获取到指数数据')

    headlines = fetch_headlines(limit=12)
    factors = build_key_factors(index_moves, headlines)

    # 新增一个交易日小节
    section_title = f"{trade_date} 美股复盘"
    doc.add_heading(section_title, level=1)
    doc.add_paragraph(f"生成时间：{now.strftime('%Y-%m-%d %H:%M:%S')}（Asia/Shanghai）")

    doc.add_heading('一、主要指数表现', level=2)
    for _, v in index_moves.items():
        arrow = '上涨' if v['pct'] >= 0 else '下跌'
        doc.add_paragraph(
            f"{v['name']}：收于 {v['close']:.2f}，较前一交易日{arrow} {abs(v['chg']):.2f} 点（{v['pct']:.2f}%）。",
            style='List Bullet'
        )

    doc.add_heading('二、新闻要点（最新交易日前后）', level=2)
    if headlines:
        for h in headlines:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"[{h['source']}] {h['title']}")
            if h['link']:
                doc.add_paragraph(h['link'])
    else:
        doc.add_paragraph('未抓取到可用新闻源。', style='List Bullet')

    doc.add_heading('三、影响走势的核心因素（研判）', level=2)
    for f in factors:
        doc.add_paragraph(f, style='List Number')

    doc.add_heading('四、次日关注清单', level=2)
    watch_items = [
        '美债收益率与美元方向是否继续走强；',
        '盘前/盘后重要公司财报与指引变化；',
        '宏观数据（通胀、就业、消费）与美联储官员表态；',
        '能源与大宗商品波动对风险偏好的传导。',
    ]
    for w in watch_items:
        doc.add_paragraph(w, style='List Bullet')

    doc.add_paragraph('----')
    doc.save(DESKTOP_DOC)
    print(DESKTOP_DOC)


if __name__ == '__main__':
    main()
